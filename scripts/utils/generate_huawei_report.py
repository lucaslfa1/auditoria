import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = docx.Document()
    
    # Estilo de Título
    title = doc.add_heading('Relatório de Diagnóstico: Falhas no Download de Gravações (AICC)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadados
    p = doc.add_paragraph()
    p.add_run('Data: 4 de maio de 2026\n').italic = True
    p.add_run('Projeto: Auditoria nstech\n').italic = True
    p.add_run('Assunto: Falha sistemática na coleta de áudio via API e OBS').italic = True
    
    # Seção 1: Contexto
    doc.add_heading('1. Contexto do Problema', level=1)
    doc.add_paragraph(
        "Nossa integração consegue listar e descobrir as ligações com sucesso. Utilizamos tanto a API CMS (/querycalls) "
        "quanto a leitura dos manifestos diários (Contact_Record em CSV) depositados no OBS. O problema ocorre "
        "exclusivamente na obtenção do arquivo de áudio."
    )
    doc.add_paragraph(
        "No último ciclo de sincronização, descobrimos 1.824 chamadas e tentamos baixar o áudio de 20 ligações "
        "utilizando 3 métodos em cascata. Todos os métodos falharam (0 sucessos, 60 tentativas totais)."
    )
    
    # Seção 2: Diagnóstico por Método
    doc.add_heading('2. Diagnóstico por Método de Download', level=1)
    
    # Método 1
    doc.add_heading('Método 1: OBS Direto (Método Primário)', level=2)
    p = doc.add_paragraph()
    p.add_run('Funcionamento: ').bold = True
    p.add_run('Leitura dos manifestos CSV na pasta Contact_Record/ do bucket OBS para confirmar o callId e o recordId. '
              'Busca direta na pasta Voice/{YYYYMMDD}/{prefixo}/.')
    
    p = doc.add_paragraph()
    p.add_run('Problema: ').bold = True
    p.add_run('O manifesto CSV confirma a chamada, mas o arquivo .V3 não é encontrado na pasta Voice/ mesmo iterando prefixos.')
    
    doc.add_paragraph('Perguntas para a Huawei:', style='List Bullet')
    doc.add_paragraph('A gravação está habilitada e configurada para o bucket obs-nstech-opentech?', style='List Bullet 2')
    doc.add_paragraph('Houve alteração na estrutura de pastas da mídia?', style='List Bullet 2')
    doc.add_paragraph('O que significa recordId vazio no Contact_Record?', style='List Bullet 2')
    
    # Método 2
    doc.add_heading('Método 2: CC-FS Binário Direto (downloadRecord)', level=2)
    p = doc.add_paragraph()
    p.add_run('Funcionamento: ').bold = True
    p.add_run('Uso do endpoint POST /CCFS/resource/ccfs/downloadRecord via API.')
    
    p = doc.add_paragraph()
    p.add_run('Problema: ').bold = True
    p.add_run('Resposta com resultCode: 0300012 ("No data found") ou timeout.')
    
    doc.add_paragraph('Perguntas para a Huawei:', style='List Bullet')
    doc.add_paragraph('O endpoint está liberado para nosso tenant?', style='List Bullet 2')
    doc.add_paragraph('Existe delay esperado entre o fim da chamada e a disponibilidade?', style='List Bullet 2')
    
    # Método 3
    doc.add_heading('Método 3: URL Pré-Assinada (getRecordFileUrlFromObs)', level=2)
    p = doc.add_paragraph()
    p.add_run('Funcionamento: ').bold = True
    p.add_run('Solicitação de URL temporária via API CC-FS.')
    
    p = doc.add_paragraph()
    p.add_run('Problema: ').bold = True
    p.add_run('Comportamento similar ao erro de "No data found".')
    
    doc.add_paragraph('Perguntas para a Huawei:', style='List Bullet')
    doc.add_paragraph('Nossa licença permite a geração de URLs pré-assinadas via API?', style='List Bullet 2')
    
    # Seção 3: Checklist
    doc.add_heading('3. Checklist de Infraestrutura (Whitelists)', level=1)
    doc.add_paragraph('Pedir para checar no console/WAF da Huawei:', style='List Bullet')
    doc.add_paragraph('IP do Proxy Nginx: 34.171.63.68', style='List Bullet 2')
    doc.add_paragraph('IP Fixo NAT Google Cloud: 35.199.111.152', style='List Bullet 2')
    doc.add_paragraph('Acesso à porta 28443 (CMS/FS) e tráfego HTTPS para o OBS.', style='List Bullet 2')
    
    # Conclusão/Dica
    doc.add_heading('Conclusão Recomendada', level=1)
    p = doc.add_paragraph()
    p.add_run("O fato de os manifestos CSV estarem presentes no OBS prova que o tráfego de rede e credenciais "
              "básicas estão funcionando. O problema parece estar na geração ou movimentação dos arquivos de áudio (.V3) "
              "pelo PABX.")
    
    output_path = 'Relatorio_Diagnostico_Huawei_Telefonia.docx'
    doc.save(output_path)
    print(f"Relatório gerado com sucesso: {output_path}")

if __name__ == "__main__":
    create_report()
