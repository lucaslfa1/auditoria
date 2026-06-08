import os
from docx import Document

def create_report():
    doc = Document()
    doc.add_heading('Relatório de Tentativas - Módulos de Ligações e Automação', 0)

    doc.add_heading('1. O Problema Inicial', level=1)
    doc.add_paragraph("A sincronização com a Huawei estava falhando ou puxando dados de forma ineficiente, gastando tokens de Inteligência Artificial com ligações de operadores que já haviam batido a meta mensal.")

    doc.add_heading('2. O que foi feito (Automação e Telefonia)', level=1)
    doc.add_paragraph("- Atualização de Rede (Proxy): Configuramos o IP e o proxy lab.teledatabrasil.com.br no banco e no DNS Override para bypassar bloqueios (WAF).")
    doc.add_paragraph("- Busca por Operador: A Huawei estava falhando ao buscar em lote na VDN, então alteramos o código para buscar as chamadas iterando operador por operador (agentId).")
    doc.add_paragraph("- Cota Mensal Separada: Implementamos a regra de baixar TODAS as ligações possíveis, mas só enviar para a Triagem e Auditoria (IA) 2 ligações mensais. As excedentes recebem a tag de 'Cota atingida' e não gastam tokens.")
    doc.add_paragraph("- Filtros de Duração: Reduzimos a exigência de duração de todas as ligações para 60 segundos (antes 90s/120s).")
    doc.add_paragraph("- Parada no Auditor (Awaiting Pair): A automação foi configurada para não enviar o resultado da IA direto para o supervisor. O status final gerado é 'awaiting_pair' para ir para 'Arquivos Salvos'.")
    doc.add_paragraph("- Frontend Config-Driven: A tela de Automação ganhou botões de controle de frequência (10 min a 2 horas) com integração direta no banco.")

    doc.add_heading('3. Por que as ligações não aparecem? (O Bloqueio Atual)', level=1)
    doc.add_paragraph("Nos testes executados em Produção, o robô encontrou até 200 ligações válidas. No entanto, quando o robô tenta efetuar o download do arquivo de áudio (.wav) na Huawei (endpoints 'downloadRecord' ou 'getRecordFileUrlFromObs'), a Huawei responde repetidamente com o Erro 0300012: 'No data found.'")
    doc.add_paragraph("Isso significa que o PABX registra que a ligação aconteceu (os metadados existem na VDN), mas NÃO ESTÁ GRAVANDO ou não está disponibilizando o arquivo físico do áudio para a API baixar. O sistema só cria o registro no painel quando o áudio é de fato entregue pela Huawei.")

    # Tentando salvar em C:\Users\Lucas\Desktop ou no D:\ (conforme pedido "deixe em d: no meu desktop")
    # Como não temos certeza de onde está o desktop do usuário, vamos salvar na raiz D:\
    output_path = r'D:\Relatorio_Tentativas_Automacao_Ligacoes.docx'
    doc.save(output_path)
    print(f"Relatório salvo com sucesso em: {output_path}")

if __name__ == '__main__':
    create_report()