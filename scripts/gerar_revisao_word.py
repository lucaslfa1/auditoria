"""Gera documento Word com as divergencias do benchmark para revisao pelo setor de auditoria."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# Estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Titulo
doc.add_heading('Revisao de Divergencias - Benchmark de Auditoria IA', level=1)
doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y")}')
doc.add_paragraph('Gerado por: Lucas Afonso - Equipe de Qualidade')
doc.add_paragraph('')

intro = doc.add_paragraph()
run = intro.add_run('Objetivo: ')
run.bold = True
intro.add_run(
    'Durante a calibracao da IA Auditora, identificamos 5 ligacoes onde o veredito da IA '
    'divergiu da classificacao da empresa. Solicitamos a revisao dessas ligacoes para '
    'alinhar os criterios de avaliacao.'
)
doc.add_paragraph('')

# Resumo
doc.add_heading('Resumo das Divergencias', level=2)
summary_table = doc.add_table(rows=1, cols=4)
summary_table.style = 'Light Grid Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = summary_table.rows[0].cells
for i, txt in enumerate(['Tipo', 'Qtd', 'Descricao', 'Impacto']):
    hdr[i].text = txt
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True

row = summary_table.add_row().cells
row[0].text = 'Falso Positivo'
row[1].text = '3'
row[2].text = 'IA aprovou ligacoes que a empresa reprovou'
row[3].text = 'Operador ruim passa como bom'

row = summary_table.add_row().cells
row[0].text = 'Falso Negativo'
row[1].text = '2'
row[2].text = 'IA reprovou ligacoes que a empresa aprovou'
row[3].text = 'Operador bom e prejudicado'

doc.add_paragraph('')


def add_colored_field(doc, label, value, color=None):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run = p.add_run(value)
    run.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_criteria_table(doc, data):
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    h = t.rows[0].cells
    for i, txt in enumerate(['Criterio', 'Status IA', 'Comentario IA']):
        h[i].text = txt
        for p in h[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for label, status, comment in data:
        row = t.add_row().cells
        row[0].text = label
        row[1].text = status
        row[2].text = comment


RED = RGBColor(255, 0, 0)
GREEN = RGBColor(0, 128, 0)

# =============================================
# SECAO 1: Falsos Positivos
# =============================================
doc.add_heading('SECAO 1: Ligacoes RUINS que a IA classificou como BOAS', level=2)
doc.add_paragraph(
    'A IA nao encontrou falhas suficientes para reprovar. '
    'Precisamos entender o que a empresa identificou de errado nessas ligacoes.'
)
doc.add_paragraph('')

# --- Caso 1: Cadastro ---
doc.add_heading('Caso 1 - Cadastro / Antecedentes', level=3)
add_colored_field(doc, 'Arquivo', 'ANTECEDENTES-agent-10362-8_10_2025_12_25_17-node01-1759937115-60427.wav')
p = doc.add_paragraph()
p.add_run('Caminho completo: ').bold = True
p.add_run(r'C:\Users\lucas.afonso\projetos\auditoria\Ligacoes\CADASTRO\RUINS')
p = doc.add_paragraph()
p.add_run('Setor: ').bold = True
p.add_run('Cadastro')
add_colored_field(doc, 'Classificacao empresa', 'RUIM', RED)
add_colored_field(doc, 'Classificacao IA', 'BOM (81.4%)', GREEN)

doc.add_paragraph('')
doc.add_paragraph('O que a IA avaliou:')

add_criteria_table(doc, [
    ('Identificacao (saudacao, nome, setor, empresa)', 'PASS', 'O operador se identificou corretamente.'),
    ('Solicitou CPF/Placa para iniciar atendimento', 'PASS', 'CPF foi solicitado para iniciar.'),
    ('Enfatizou bloqueio/cadastro negativado', 'PASS', 'Informou sobre o bloqueio.'),
    ('Informou inquerito/processo/apontamento', 'PASS', 'Forneceu informacoes sobre o processo.'),
    ('Informou estado/justica federal', 'FAIL', 'Nao informou qual o estado ou justica federal.'),
    ('Informou documento necessario', 'PASS', 'Documentos requeridos foram informados.'),
    ('Despedida cordial', 'PASS', 'Despedida adequada.'),
    ('Silencio prolongado (>45s)', 'FAIL', 'Houve silencio prolongado superior a 45 segundos.'),
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Perguntas para o setor de auditoria:').bold = True
doc.add_paragraph('1. Qual foi o motivo PRINCIPAL da reprovacao desta ligacao?', style='List Number')
doc.add_paragraph('2. O silencio prolongado (>45s) foi o fator determinante? Ou houve outro problema?', style='List Number')
doc.add_paragraph('3. O operador deixou de informar algo critico que nao esta nos criterios 4.2.1?', style='List Number')
doc.add_paragraph('4. O tom de voz ou postura do operador foi inadequado (algo que a transcricao nao captura)?', style='List Number')

doc.add_page_break()

# --- Caso 2: Unilever Tratativa ---
doc.add_heading('Caso 2 - Unilever / Atuacao Tratativa', level=3)
add_colored_field(doc, 'Arquivo', 'ATUACAO TRATATIVA-CLIENTE-05-09-2025_12-57-32_11353_11953217745.mp3')
p = doc.add_paragraph()
p.add_run('Caminho completo: ').bold = True
p.add_run(r'C:\Users\lucas.afonso\projetos\auditoria\Ligacoes\UNILEVER\RUINS')
p = doc.add_paragraph()
p.add_run('Setor: ').bold = True
p.add_run('Logistica Unilever')
add_colored_field(doc, 'Classificacao empresa', 'RUIM', RED)
add_colored_field(doc, 'Classificacao IA', 'BOM (86.1%)', GREEN)

doc.add_paragraph('')
doc.add_paragraph('O que a IA avaliou:')

add_criteria_table(doc, [
    ('Identificacao (saudacao, nome, setor, empresa)', 'PARTIAL', 'Identificou nome e setor, mas informou empresa como Pentech em vez de Opentech.'),
    ('Confirmou com quem esta falando', 'PASS', 'Confirmou o interlocutor.'),
    ('Informou motivo do contato', 'PASS', 'Motivo do contato informado.'),
    ('Informou nome do cliente', 'PASS', 'Nome informado corretamente.'),
    ('Informou endereco do cliente', 'PASS', 'Endereco informado.'),
    ('Informou codigo do cliente', 'PASS', 'Codigo informado.'),
    ('Informou motivo da devolucao', 'PASS', 'Motivo da devolucao informado.'),
    ('Informou quantidade de caixas', 'PASS', 'Quantidade informada.'),
    ('Informou tempo de espera', 'FAIL', 'Nao informou o tempo de espera.'),
    ('Acao resultante registrada', 'PASS', 'Acao registrada.'),
    ('Despedida cordial', 'PARTIAL', 'Encerrou a ligacao abruptamente.'),
    ('Entonacao e cordialidade', 'PARTIAL', 'Houve momentos de confusao na conducao.'),
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Perguntas para o setor de auditoria:').bold = True
doc.add_paragraph('1. Qual foi o motivo PRINCIPAL da reprovacao desta ligacao?', style='List Number')
doc.add_paragraph('2. O fato de ter informado "Pentech" em vez de "Opentech" e motivo de reprovacao?', style='List Number')
doc.add_paragraph('3. A falta do tempo de espera e a despedida abrupta foram os fatores determinantes?', style='List Number')
doc.add_paragraph('4. Existe algum criterio nao documentado no 4.3.3 que o operador descumpriu?', style='List Number')
doc.add_paragraph('5. O tom de voz ou postura do operador foi inadequado?', style='List Number')

doc.add_paragraph('')

# --- Caso 3: Unilever Devolução ---
doc.add_heading('Caso 3 - Unilever / Devolucao', level=3)
add_colored_field(doc, 'Arquivo', 'DEVOLUCAO-CLIENTE-22-08-2025_11-33-33_11231_11976013982.mp3')
p = doc.add_paragraph()
p.add_run('Caminho completo: ').bold = True
p.add_run(r'C:\Users\lucas.afonso\projetos\auditoria\Ligacoes\UNILEVER\RUINS')
p = doc.add_paragraph()
p.add_run('Setor: ').bold = True
p.add_run('Logistica Unilever')
add_colored_field(doc, 'Classificacao empresa', 'RUIM', RED)
add_colored_field(doc, 'Classificacao IA', 'BOM (80.6%)', GREEN)

doc.add_paragraph('')
doc.add_paragraph('O que a IA avaliou:')

add_criteria_table(doc, [
    ('Identificacao (saudacao, nome, setor, empresa)', 'PARTIAL', 'Saudacao e setor informados, mas empresa nao mencionada claramente.'),
    ('Confirmou com quem esta falando', 'PASS', 'Confirmou o interlocutor.'),
    ('Informou devolucao confirmada e proximo passo', 'PASS', 'Informou sobre a devolucao.'),
    ('Informou nome do cliente', 'PASS', 'Nome informado.'),
    ('Informou endereco do cliente', 'PASS', 'Endereco informado.'),
    ('Informou codigo do cliente', 'PASS', 'Codigo informado.'),
    ('Confirmou quantidade de caixas', 'PASS', 'Quantidade confirmada.'),
    ('Acao resultante registrada', 'FAIL', 'Nao houve registro explicito da acao resultante.'),
    ('Despedida cordial', 'PARTIAL', 'Despedida sem mensagem padrao como Tenha um bom dia.'),
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Perguntas para o setor de auditoria:').bold = True
doc.add_paragraph('1. Qual foi o motivo PRINCIPAL da reprovacao desta ligacao?', style='List Number')
doc.add_paragraph('2. A falta de registro da acao resultante foi determinante?', style='List Number')
doc.add_paragraph('3. Houve algum erro de informacao (nome, endereco, codigo incorretos)?', style='List Number')
doc.add_paragraph('4. O operador deveria ter feito algo que nao consta nos criterios 4.3.1?', style='List Number')

doc.add_page_break()

# =============================================
# SECAO 2: Falsos Negativos
# =============================================
doc.add_heading('SECAO 2: Ligacoes BOAS que a IA classificou como RUINS', level=2)
doc.add_paragraph(
    'A IA encontrou falhas que resultaram em reprovacao. '
    'Precisamos confirmar se esses criterios realmente se aplicam ao setor de Logistica.'
)
doc.add_paragraph('')

# --- Caso 4: Logística Atraso ---
doc.add_heading('Caso 4 - Logistica / Atraso Motorista', level=3)
add_colored_field(doc, 'Arquivo', 'ATRASO-MOTORISTA-20251230173926115_Danilo_Alves_Logistica_Voz.wav')
p = doc.add_paragraph()
p.add_run('Caminho completo: ').bold = True
p.add_run(r'C:\Users\lucas.afonso\projetos\auditoria\Ligacoes\LOGISTICA\BOAS')
p = doc.add_paragraph()
p.add_run('Setor: ').bold = True
p.add_run('Logistica Opentech')
add_colored_field(doc, 'Classificacao empresa', 'BOA', GREEN)
add_colored_field(doc, 'Classificacao IA', 'RUIM (53.8%)', RED)

doc.add_paragraph('')
doc.add_paragraph('O que a IA reprovou:')

add_criteria_table(doc, [
    ('Identificacao (saudacao, nome, setor, empresa)', 'PASS', 'Identificacao correta.'),
    ('Confirmou com quem esta falando', 'PASS', 'Interlocutor confirmado.'),
    ('Informou motivo do contato', 'PASS', 'Motivo informado.'),
    ('Confirmou localizacao atual do motorista', 'FAIL', 'Nao confirmou a localizacao atual do motorista.'),
    ('Orientou forcar posicionamento do rastreador', 'FAIL', 'Nao orientou sobre posicionamento do rastreador.'),
    ('Identificou motivo da perda de sinal/atraso', 'PASS', 'Motivo identificado.'),
    ('Informou riscos operacionais', 'FAIL', 'Nao informou riscos operacionais.'),
    ('Despedida cordial', 'PASS', 'Despedida cordial.'),
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Perguntas para o setor de auditoria:').bold = True
doc.add_paragraph(
    '1. No setor de Logistica, o operador precisa "forcar posicionamento do rastreador" em alertas '
    'de atraso? Ou isso so se aplica ao BAS/UTI?', style='List Number')
doc.add_paragraph(
    '2. O operador de Logistica precisa informar "riscos operacionais e de seguro" ao motorista? '
    'Ou basta confirmar a situacao e registrar?', style='List Number')
doc.add_paragraph(
    '3. Quais sao os criterios REAIS usados para avaliar ligacoes de Atraso na Logistica? '
    'Os criterios 4.4.x do documento contemplam esse tipo de alerta?', style='List Number')
doc.add_paragraph('4. Essa ligacao foi considerada BOA por quais motivos especificos?', style='List Number')

doc.add_paragraph('')

# --- Caso 5: Logística Desvio ---
doc.add_heading('Caso 5 - Logistica / Desvio de Rota Motorista', level=3)
add_colored_field(doc, 'Arquivo', 'DESVIO-MOTORISTA-agent-11218-19_11_2025_19_57_50-node01-1763593067-198710.wav')
p = doc.add_paragraph()
p.add_run('Caminho completo: ').bold = True
p.add_run(r'C:\Users\lucas.afonso\projetos\auditoria\Ligacoes\LOGISTICA\BOAS')
p = doc.add_paragraph()
p.add_run('Setor: ').bold = True
p.add_run('Logistica Opentech')
add_colored_field(doc, 'Classificacao empresa', 'BOA', GREEN)
add_colored_field(doc, 'Classificacao IA', 'RUIM (69.2%)', RED)

doc.add_paragraph('')
doc.add_paragraph('O que a IA reprovou:')

add_criteria_table(doc, [
    ('Identificacao (saudacao, nome, setor, empresa)', 'PASS', 'Identificacao correta.'),
    ('Confirmou com quem esta falando', 'PASS', 'Interlocutor confirmado.'),
    ('Informou motivo do contato', 'PASS', 'Motivo informado.'),
    ('Confirmou motivo do desvio', 'PASS', 'Motivo do desvio confirmado.'),
    ('Verificou plano de viagem', 'PASS', 'Plano de viagem verificado.'),
    ('Orientou retornar a rota', 'FAIL', 'Nao orientou o motorista a retornar para a rota.'),
    ('Informou riscos do desvio', 'FAIL', 'Nao informou os riscos operacionais do desvio.'),
    ('Despedida cordial', 'PASS', 'Despedida cordial.'),
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Perguntas para o setor de auditoria:').bold = True
doc.add_paragraph(
    '1. No setor de Logistica, o operador DEVE orientar o motorista a retornar a rota? '
    'Ou basta registrar a situacao?', style='List Number')
doc.add_paragraph(
    '2. O operador de Logistica precisa informar riscos operacionais em caso de desvio? '
    'Ou isso e exclusivo do BAS/UTI?', style='List Number')
doc.add_paragraph(
    '3. Quais sao os criterios REAIS para Desvio de Rota na Logistica? '
    'Existem criterios documentados alem do 4.4.x?', style='List Number')
doc.add_paragraph('4. Essa ligacao foi considerada BOA por quais motivos especificos?', style='List Number')

doc.add_page_break()

# =============================================
# SECAO 3: Perguntas Gerais
# =============================================
doc.add_heading('Perguntas Gerais para Alinhamento', level=2)

doc.add_paragraph(
    '1. Qual e a nota minima (percentual) para uma ligacao ser considerada BOA? '
    'Atualmente a IA usa 70%. A empresa usa o mesmo limiar?', style='List Number')
doc.add_paragraph(
    '2. Existem criterios de auditoria para o setor de Logistica alem dos documentados na secao 4.4 '
    '(que so cobre Estadia)? Precisamos dos criterios completos para: Atraso, Desvio de Rota, '
    'Parada Indevida, Posicao em Atraso e Temperatura.', style='List Number')
doc.add_paragraph(
    '3. Ha situacoes em que uma ligacao e reprovada por motivos NAO previstos nos criterios formais? '
    '(ex: tom de voz, postura, erro de sistema)', style='List Number')
doc.add_paragraph(
    '4. A identificacao do operador precisa incluir TODOS os 4 itens (saudacao + nome + setor + empresa) '
    'para ser PASS? Ou saudacao + nome ja e suficiente?', style='List Number')
doc.add_paragraph(
    '5. O criterio de "registro no sistema" e "qualificacao do atendimento" pode ser verificado '
    'apenas por audio? Ou depende de validacao no sistema?', style='List Number')

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('Agradecemos o retorno para que possamos calibrar a IA Auditora com maxima precisao.').italic = True

output_path = r'C:\Users\lucas.afonso\projetos\auditoria\logs\Revisao_Divergencias_Benchmark_IA.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
