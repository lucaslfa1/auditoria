import os
import subprocess
import sys

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

ensure_docx()

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_report():
    doc = Document()
    
    title = doc.add_heading('Relatório Evolutivo: Métodos de Autenticação e Integração Huawei AICC', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('Data: 27 de Abril de 2026').bold = True
    doc.add_paragraph('Projeto: Auditoria').bold = True
    doc.add_paragraph('Diretório Base: C:\\Users\\lucas.afonso\\projetos\\auditoria').bold = True
    
    doc.add_paragraph('---')
    
    doc.add_heading('1. Introdução', level=1)
    doc.add_paragraph('O objetivo deste documento é apresentar o histórico cronológico de todas as abordagens de autenticação testadas e implementadas para a integração do módulo de Telefonia com a nuvem da Huawei (AICC). Este relatório serve como documentação definitiva para embasar revisões de arquitetura, auditorias de segurança e troubleshooting futuro.')
    
    doc.add_heading('2. Fase 1: Tentativa Inicial de Acesso Direto (Falha 403)', level=1)
    doc.add_paragraph('Período: Configuração inicial do projeto. | Status: Descartada.').italic = True
    doc.add_paragraph('- Abordagem: Tentativa de comunicação nativa direta com a API da Huawei (portal brazilsaas.aicccloud.com), solicitando um token de acesso através do endpoint /apigovernance/api/oauth/tokenByAkSk.', style='List Bullet')
    doc.add_paragraph('- Credenciais Utilizadas: As chaves de ambiente iniciais fornecidas à equipe (Ex: HUAWEI_AK iniciando com 97501...).', style='List Bullet')
    doc.add_paragraph('- Resultado: A API retornava consistentemente o erro HTTP 403 Forbidden ou falhas de bloqueio via Nginx.', style='List Bullet')
    doc.add_paragraph('- Diagnóstico da Época: Assumiu-se que o problema era um bloqueio de infraestrutura no WAF (Web Application Firewall) da Huawei, motivado pela ausência dos IPs do Google Cloud na whitelist do AICC.', style='List Bullet')
    
    doc.add_heading('3. Fase 2: Contingência via Proxy Teledata', level=1)
    doc.add_paragraph('Período: Implementação do contorno até a manhã de 27 de Abril. | Status: Descontinuado (Legado).').italic = True
    doc.add_paragraph('- Abordagem: Diante do bloqueio direto, a arquitetura foi pivotada para o uso de um proxy parceiro (Teledata). Neste modelo, nosso sistema não autenticava diretamente na Huawei, mas delegava a assinatura dos pacotes (SDK-HMAC-SHA256) para um script PHP da Teledata (c2Authorization.php).', style='List Bullet')
    doc.add_paragraph('- Configurações do Ambiente:', style='List Bullet')
    doc.add_paragraph('  - HUAWEI_AUTH_MODE=proxy')
    doc.add_paragraph('  - HUAWEI_PROXY_URL=https://lab.teledatabrasil.com.br/aicc/auth/c2Authorization.php (ou IP 163.176.162.83)')
    doc.add_paragraph('  - HUAWEI_APP_KEY=97501... (Chave revelada posteriormente como exclusiva deste proxy).')
    doc.add_paragraph('- Resultado: A integração funcionou e garantiu a operação. No entanto, introduziu-se um "débito técnico" e um ponto único de falha (Single Point of Failure) externo à arquitetura do Google Cloud, adicionando latência e dependência da estabilidade da Teledata.', style='List Bullet')
    
    doc.add_heading('4. Fase 3: Homologação da Whitelist e Diagnóstico Definitivo', level=1)
    doc.add_paragraph('Período: 27 de Abril (Prova de Conceito). | Status: Homologado.').italic = True
    doc.add_paragraph('- Abordagem: A Huawei confirmou a inclusão do IP de saída do Cloud NAT (35.199.111.152) em sua whitelist. Foi criado um script isolado (tmp_test_huawei_oauth.py) executado diretamente no Cloud Run para validar o fim do bloqueio.', style='List Bullet')
    doc.add_paragraph('- A Descoberta (A Causa Real do 403): Ao testar o IP liberado com as credenciais antigas (97501...), o erro 403 persistiu. Através de engenharia reversa na Coleção Postman utilizada pelo cliente, descobriu-se que a chave 97501... pertencia exclusivamente ao Proxy da Teledata. A chave oficial para acesso direto à nuvem da Huawei exigia um formato UUID.', style='List Bullet')
    doc.add_paragraph('- Teste de Sucesso: Ao injetar a chave correta encontrada no Postman, o sistema furou o bloqueio imediatamente e retornou HTTP 200 OK (Token gerado com sucesso).', style='List Bullet')
    
    doc.add_heading('5. Fase 4: Refatoração para OAuth Direto (Arquitetura Atual)', level=1)
    doc.add_paragraph('Período: 27 de Abril (Ativo). | Status: Em Produção (Recomendado).').italic = True
    doc.add_paragraph('- Abordagem: O cliente HTTP do backend (huawei_client.py) foi profundamente reescrito (via Codex) para adotar a comunicação direta. A aplicação realiza chamadas seguras para o AICC e possui um sistema robusto de Cache em Memória, que armazena o AccessToken por até 3300 segundos, evitando Rate Limiting por excesso de logins. O proxy da Teledata foi abandonado.', style='List Bullet')
    doc.add_paragraph('- Configuração Ativa no Google Cloud Run (Regiões: southamerica-east1 e us-central1):', style='List Bullet')
    doc.add_paragraph('  - HUAWEI_AUTH_MODE=oauth_direct')
    doc.add_paragraph('  - HUAWEI_DIRECT_APP_KEY=<HUAWEI_DIRECT_APP_KEY>')
    doc.add_paragraph('  - HUAWEI_DIRECT_APP_SECRET=<HUAWEI_DIRECT_APP_SECRET>')
    doc.add_paragraph('  - HUAWEI_TENANT_SPACE_ID=202509298231')
    doc.add_paragraph('- Headers Injetados Dinamicamente:', style='List Bullet')
    doc.add_paragraph('  - Authorization: Bearer <AccessToken>')
    doc.add_paragraph('  - X-APP-Key: <HUAWEI_DIRECT_APP_KEY>')
    doc.add_paragraph('  - X-TenantSpaceID: 202509298231')
    
    doc.add_heading('6. Conclusão para Revisão Externa', level=1)
    doc.add_paragraph('A evolução da autenticação culminou na remoção completa do Middleman (Teledata), resultando em uma comunicação Ponto-a-Ponto mais limpa, rápida e em estrita conformidade com as diretrizes OAuth da Huawei.')
    doc.add_paragraph('A arquitetura atual (Fase 4) é a definitiva.')
    doc.add_paragraph('Recomendação para a Próxima Etapa: Solicitar as credenciais do bucket OBS da Huawei (HUAWEI_OBS_AK, HUAWEI_OBS_SK, HUAWEI_OBS_BUCKET, HUAWEI_OBS_ENDPOINT) para viabilizar o download de contingência de gravações ("Fallback"), visto que algumas APIs de extração nativa estão reportando o erro Huawei 0300012 No data found.').bold = True
    
    os.makedirs('relatorios_atividades', exist_ok=True)
    file_path = os.path.join('relatorios_atividades', 'historico_autenticacao_huawei.docx')
    doc.save(file_path)
    print(f"Documento Word gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    create_word_report()
