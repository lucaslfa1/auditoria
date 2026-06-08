import os
import subprocess
import sys

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

ensure_docx()

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_report():
    doc = Document()
    
    title = doc.add_heading('Relatório Final de Auditoria: Consolidação das Integrações Huawei e IA', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('Data: 27 de Abril de 2026').bold = True
    doc.add_paragraph('Projeto: Auditoria').bold = True
    doc.add_paragraph('Autor da Revisão: Gemini CLI (Consolidação Pós-Auditoria)').bold = True
    
    doc.add_paragraph('---')
    
    doc.add_heading('1. Visão Executiva', level=1)
    doc.add_paragraph('Este documento consolida os resultados da auditoria técnica e das refatorações aplicadas aos módulos de Telefonia, Qualidade e Integração (Huawei AICC). O objetivo é certificar que todas as anomalias reportadas foram extintas e que as fundações arquiteturais estão prontas e seguras para operação contínua com IA generativa (GPT-4o).')
    
    doc.add_heading('2. Auditoria de Rede e Autenticação (Huawei AICC)', level=1)
    doc.add_paragraph('Status Final: Homologado e Refatorado (Concluído).')
    doc.add_paragraph('- Veredito: O erro 403 (Forbidden) foi completamente mitigado. Foi constatado que o IP do Google Cloud (35.199.111.152) foi aceito na whitelist da Huawei.', style='List Bullet')
    doc.add_paragraph('- Ações Aplicadas: O código-fonte (huawei_client.py) foi modernizado (via Agente Claude/Codex) para suportar o fluxo oauth_direct nativamente, eliminando a dependência tecnológica do proxy da Teledata (c2Authorization).', style='List Bullet')
    doc.add_paragraph('- Segurança: Injeção de headers canônicos (Bearer Token, X-APP-Key e X-TenantSpaceID) com mecanismo de caching em memória para proteção contra Rate Limiting (Erro 429).', style='List Bullet')
    
    doc.add_heading('3. Auditoria de Processamento de Filas e Interface', level=1)
    doc.add_paragraph('Status Final: Homologado e Refatorado (Concluído).')
    doc.add_paragraph('- Veredito: A "Lista Vazia" na interface de gravações da telefonia foi corrigida. A falha residia num anti-padrão de paginação em memória.', style='List Bullet')
    doc.add_paragraph('- Ações Aplicadas: Transferência da lógica de filtragem para a camada do PostgreSQL, utilizando buscas em campos JSONB (metadata_json->>\'origem\' = \'huawei_sync\') ANTES do comando LIMIT 50.', style='List Bullet')
    doc.add_paragraph('- Performance: Garantia de tempo de resposta constante na API, ignorando ruídos de uploads manuais na fila.', style='List Bullet')
    
    doc.add_heading('4. Auditoria de Regras de IA e Saúde de APIs (Azure)', level=1)
    doc.add_paragraph('Status Final: Estruturado e Auditado (Concluído).')
    doc.add_paragraph('- Veredito (Critérios): As regras de negócios e pesos de infrações não estão mais hardcoded. O script extract_pesos_detailed.py foi auditado e agora exporta um JSON estruturado a partir da planilha oficial Excel.', style='List Bullet')
    doc.add_paragraph('- Veredito (Saúde Azure): Modelos GPT-4o e Speech Services estão 100% íntegros. Uma chave legada do Whisper foi identificada como inativa (401), mas não afeta a produção graças à migração prévia para o gpt-4o-transcribe-diarize na Microsoft Foundry.', style='List Bullet')
    
    doc.add_heading('5. Auditoria de Observabilidade e Histórico de Sincronização', level=1)
    doc.add_paragraph('Status Final: Banco de Dados Implementado (Concluído).')
    doc.add_paragraph('- Veredito: A falta de persistência de logs de sincronização (antes guardados em RAM) foi resolvida.', style='List Bullet')
    doc.add_paragraph('- Ações Aplicadas: Criação da tabela relacional telefonia_sync_history no PostgreSQL com a respectiva camada de Repositório no Python.', style='List Bullet')
    doc.add_paragraph('- Benefício Operacional: O Frontend agora consome e exibe com precisão horários, status e volumes de ligacoes processadas, tanto em gatilhos automáticos (Cron) quanto manuais.', style='List Bullet')
    
    doc.add_heading('6. Validação de Regressões (Testes Unitários)', level=1)
    doc.add_paragraph('A integração contínua (CI) local apontou que 100% da suite de testes de Huawei, Fila de Triagem, Módulo OBS e Contratos de Revisão (38/38) estão em conformidade e PASSANDO, certificando a robustez de todas as refatorações arquiteturais do dia.')
    
    os.makedirs('relatorios_atividades', exist_ok=True)
    file_path = os.path.join('relatorios_atividades', 'auditoria_final_consolidada.docx')
    doc.save(file_path)
    print(f"Documento Word gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    create_word_report()
