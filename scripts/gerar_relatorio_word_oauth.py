import os
import subprocess
import sys

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

ensure_docx()

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_report():
    doc = Document()
    
    # Título principal
    title = doc.add_heading('Relatório Técnico: Refatoração da Autenticação Huawei AICC para Acesso Direto (OAuth 2.0)', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadados
    doc.add_paragraph('Data: 27 de Abril de 2026').bold = True
    doc.add_paragraph('Projeto: Auditoria').bold = True
    doc.add_paragraph('Diretório Base: C:\\Users\\lucas.afonso\\projetos\\auditoria').bold = True
    
    doc.add_paragraph('---')
    
    # 1. Contexto e Objetivos
    doc.add_heading('1. Contexto e Objetivos', level=1)
    doc.add_paragraph('Dando continuidade à homologação do IP Estático 35.199.111.152 na whitelist da Huawei (comprovado via Prova de Conceito isolada no Cloud Run), este relatório detalha a refatoração definitiva do cliente de integração da telefonia (huawei_client.py).')
    doc.add_paragraph('O objetivo principal foi desativar a dependência do Proxy C2 (Teledata) e implementar nativamente o fluxo de autenticação direta oauth_direct, garantindo alta performance, segurança e alinhamento total com as especificações testadas via Postman.')
    
    # 2. Atividade Principal
    doc.add_heading('2. Atividade Principal: Implementação do Fluxo oauth_direct', level=1)
    
    doc.add_heading('2.1. Arquitetura de Autenticação e Cache de Tokens', level=2)
    doc.add_paragraph('A classe HuaweiAICCClient foi profundamente reestruturada para suportar múltiplos modos de autenticação de forma transparente:')
    doc.add_paragraph('- Novo Modo oauth_direct (Alias: token): Substitui a lógica antiga de assinatura HMAC e proxy. O sistema agora faz uma requisição inicial POST para tokenByAkSk utilizando as credenciais diretas do AICC.', style='List Bullet')
    doc.add_paragraph('- Gerenciamento de Estado (Cache): Para evitar sobrecarga de requisições de autenticação e risco de rate limiting por parte da Huawei (Erro 429), foi implementado um mecanismo de cache em memória do AccessToken.', style='List Bullet')
    doc.add_paragraph('  - O token é armazenado com base no tempo de vida (campo expiresIn), que possui um default de 3300 segundos.')
    doc.add_paragraph('  - Foi introduzido um buffer de segurança de 60 segundos, garantindo que o cliente solicite proativamente um novo token antes do atual expirar.')
    
    doc.add_heading('2.2. Separação Estrita de Credenciais', level=2)
    doc.add_paragraph('Para mitigar os falsos-positivos de bloqueio (Erro 403) causados por colisão de credenciais (onde chaves do proxy Teledata eram enviadas indevidamente para a porta oficial da Huawei), o sistema de injeção de dependências foi atualizado:')
    doc.add_paragraph('- Novas Variáveis Introduzidas:', style='List Bullet')
    doc.add_paragraph('  - HUAWEI_DIRECT_APP_KEY')
    doc.add_paragraph('  - HUAWEI_DIRECT_APP_SECRET')
    doc.add_paragraph('  - HUAWEI_TENANT_SPACE_ID')
    doc.add_paragraph('- Lógica de Fallback: O código prioriza as chaves DIRECT_APP_*. Caso não existam, o sistema recorre às antigas HUAWEI_APP_KEY/SECRET para manter compatibilidade retroativa durante a transição.', style='List Bullet')
    doc.add_paragraph('- Otimização de Base URL: A auth_base_url agora é resolvida por precedência:', style='List Bullet')
    doc.add_paragraph('  1. Override explícito.')
    doc.add_paragraph('  2. Variável de ambiente (HUAWEI_AUTH_BASE_URL ou HUAWEI_PORTAL_URL).')
    doc.add_paragraph('  3. Derivação inteligente do HUAWEI_CMS_URL, onde a porta de dados (:28443) é automaticamente suprimida para construir a URL do portal de governança.')
    
    doc.add_heading('2.3. Formatação Canônica de Cabeçalhos (Headers)', level=2)
    doc.add_paragraph('Para respeitar estritamente o WAF e o API Gateway da Huawei, a construção de cabeçalhos no modo oauth_direct foi ajustada:')
    doc.add_paragraph('- A propriedade Authorization passou a receber o valor prefixado por Bearer <token>.', style='List Bullet')
    doc.add_paragraph('- O cabeçalho X-APP-Key teve sua grafia forçada (uppercase), conforme exigência não-documentada descoberta via reverse engineering da coleção Postman.', style='List Bullet')
    doc.add_paragraph('- Inclusão sistemática do X-TenantSpaceID (Ex: 202509298231), crucial para roteamento multitenant no AICC.', style='List Bullet')
    
    doc.add_paragraph('---')
    
    # 3. Atividade Secundária
    doc.add_heading('3. Atividade Secundária: Resolução de Débitos Técnicos e Testes', level=1)
    
    doc.add_heading('3.1. Restauração e Expansão da Suíte de Testes', level=2)
    doc.add_paragraph('O ecossistema de testes unitários encontrava-se parcialmente degradado devido aos commits recentes da funcionalidade "Auditar Instantaneamente".')
    doc.add_paragraph('- Fila de Revisão (test_review_queue_contract.py): Foram corrigidos 2 testes classificados como stale (linhas 316 e 431). O erro decorria da ausência do assert para o novo status awaiting_pair, injetado recentemente.', style='List Bullet')
    doc.add_paragraph('- Cobertura do Cliente Huawei (test_huawei_client.py): Foram desenvolvidos 9 novos casos de teste dedicados à nova arquitetura, validando rigorosamente:', style='List Bullet')
    doc.add_paragraph('  - Operação dos aliases de autenticação (token vs oauth_direct).')
    doc.add_paragraph('  - Lógica de derivação e limpeza de portas do auth_base_url.')
    doc.add_paragraph('  - Resiliência do fallback de credenciais.')
    doc.add_paragraph('  - Validade do formato dos cabeçalhos (Bearer + Tenant).')
    doc.add_paragraph('  - Funcionamento algorítmico do Cache Hit e do ciclo de Refetch pós-expiração.')
    
    doc.add_heading('3.2. Resultado da Integração Contínua', level=2)
    doc.add_paragraph('A execução combinada das suítes de teste (Huawei Client, Huawei Sync, OBS Client e Fila de Revisão) resultou em 38 testes aprovados (38/38 PASSED), assegurando que as novas implementações não causaram regressões na cadeia de automação.')
    doc.add_paragraph('O arquivo .env.example e a documentação interna (docs/huawei/README.md) foram integralmente atualizados para refletir a nova topologia.')
    
    doc.add_paragraph('---')
    
    # 4. Plano de Ação
    doc.add_heading('4. Plano de Ação: Implantação e Virada de Chave (Go-Live)', level=1)
    doc.add_paragraph('O código está pronto e maduro. A ativação do fluxo direto na nuvem depende estritamente das seguintes variáveis de ambiente na aba "Variables & Secrets" do Google Cloud Run (revisão auditoria-nstech):')
    
    code_paragraph = doc.add_paragraph()
    code_paragraph.add_run('# Modo de Autenticação: Migração de Proxy para Acesso Direto\n').font.name = 'Courier New'
    code_paragraph.add_run('HUAWEI_AUTH_MODE=oauth_direct\n\n').font.name = 'Courier New'
    code_paragraph.add_run('# Chaves Oficiais (AICC OAuth)\n').font.name = 'Courier New'
    code_paragraph.add_run('HUAWEI_DIRECT_APP_KEY=<HUAWEI_DIRECT_APP_KEY>\n').font.name = 'Courier New'
    code_paragraph.add_run('HUAWEI_DIRECT_APP_SECRET=<HUAWEI_DIRECT_APP_SECRET>\n').font.name = 'Courier New'
    code_paragraph.add_run('HUAWEI_TENANT_SPACE_ID=202509298231\n').font.name = 'Courier New'
    
    doc.add_paragraph('(Nota de infraestrutura: Assim que o Cloud Run recarregar com as variáveis acima, o sistema descartará automaticamente o tráfego em direção ao script PHP da Teledata, passando a utilizar o IP 35.199.111.152 contra o portal nativo da Huawei).').italic = True
    
    # Salvar o documento
    os.makedirs('relatorios_atividades', exist_ok=True)
    file_path = os.path.join('relatorios_atividades', 'refatoracao_autenticacao_huawei.docx')
    doc.save(file_path)
    print(f"Documento Word gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    create_word_report()