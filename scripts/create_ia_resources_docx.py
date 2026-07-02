import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Define a cor de fundo de uma celula de tabela."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def apply_text_styling(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=None):
    """Aplica formatacao a um segmento de texto."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def create_auditoria_docx(filepath):
    doc = docx.Document()
    
    # Titulo Principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("MAPA DE RECURSOS DE INTELIGÊNCIA ARTIFICIAL\nSISTEMA DE AUDITORIA NSTECH")
    apply_text_styling(title_run, size_pt=16, bold=True, color_rgb=RGBColor(0, 32, 96))
    
    # Subtitulo
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Relatório Executivo de Tecnologias, Aplicações e Benefícios de Negócio\n")
    apply_text_styling(sub_run, size_pt=11, italic=True, color_rgb=RGBColor(100, 100, 100))
    
    # Introducao
    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run(
        "Este documento descreve as tecnologias de Inteligência Artificial (IA) integradas ao Sistema de Auditoria. "
        "O sistema foi desenhado para automatizar o controle de qualidade dos atendimentos de telefonia e chat por meio "
        "de transcrição inteligente, análise de sentimento e avaliação contra critérios de conformidade."
    )
    apply_text_styling(intro_run, size_pt=11)
    
    # Secao 1: Recursos de IA
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Recursos de Inteligência Artificial e Atuação")
    apply_text_styling(h1_run, size_pt=14, bold=True, color_rgb=RGBColor(0, 32, 96))
    
    recursos = [
        {
            "nome": "Robô de Triagem Inicial de Ligações (Triagem Inteligente)",
            "atuacao": "Processo de Sincronização e Coleta de Ligações (Huawei Sync).",
            "descricao": "Durante a importação diária de chamadas da telefonia, a IA realiza uma leitura rápida do contexto para confirmar se a ligação é de fato útil ou se representa ruído de central (enganos, quedas instantâneas).",
            "beneficio": "Redução acentuada nos custos operacionais ao evitar que o sistema processe e gaste recursos de transcrição com ligações irrelevantes."
        },
        {
            "nome": "Transcrição de Voz com Redundância Automática (Speech-to-Text)",
            "atuacao": "Processamento e entrada de áudios de ligações.",
            "descricao": "IA que converte o áudio falado da ligação em texto estruturado em forma de diálogo (separando atendente e cliente) e indicando os minutos de cada fala. O sistema conta com 4 níveis de motores de IA em cascata (caso um apresente falha de conexão ou ruído, o próximo assume automaticamente).",
            "beneficio": "Segurança operacional e estabilidade, garantindo alta fidelidade do texto extraído mesmo sob conexões de telefonia comprimidas ou de baixa qualidade."
        },
        {
            "nome": "Avaliação de Notas e Critérios de Qualidade (IA Judge)",
            "atuacao": "Motor de cálculo e pontuação de auditorias.",
            "descricao": "A IA analisa o diálogo transcrito e o pontua de acordo com o Catálogo Oficial de Critérios (ex: saudação inicial, solicitação do CPF de segurança, cordialidade, encerramento). Ela aplica penalidades e inclusive regras de 'nota zero' para falhas críticas definidas pela gerência.",
            "beneficio": "Padronização absoluta das avaliações, redução do tempo de auditoria de horas para segundos e eliminação total de subjetividade humana na nota."
        },
        {
            "nome": "Auditoria de Atendimento Escrito (Leitor Documental de Chats)",
            "atuacao": "Módulo de Auditoria Documental manual (Upload de PDFs).",
            "descricao": "Processador inteligente que lê arquivos PDF extraídos de ferramentas de chat de atendimento. Ele converte a conversa escrita, separa o diálogo de ambas as partes e aciona a mesma esteira de nota da IA usada para ligações de telefone.",
            "beneficio": "Unificação da governança de qualidade. O auditor monitora ligações faladas e chats escritos em uma única tela sob os mesmos critérios."
        },
        {
            "nome": "Módulo de Calibração e Alinhamento de Critérios (Dataset Gabarito)",
            "atuacao": "Painel Administrativo de Configuração de IA.",
            "descricao": "Permite que os auditores salvem auditorias reais avaliadas por humanos como 'gabarito'. A IA do sistema realiza testes internos de forma constante comparando seu resultado com estes gabaritos para se auto-corrigir.",
            "beneficio": "Garantia de melhoria contínua da IA e alinhamento com a equipe de qualidade, evitando desvios ou notas fora do padrão operacional desejado."
        }
    ]
    
    for r in recursos:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        
        name_run = p.add_run(f"• {r['nome']}\n")
        apply_text_styling(name_run, size_pt=12, bold=True, color_rgb=RGBColor(56, 189, 248))
        
        act_lbl = p.add_run("  Onde atua: ")
        apply_text_styling(act_lbl, size_pt=10.5, bold=True)
        act_val = p.add_run(f"{r['atuacao']}\n")
        apply_text_styling(act_val, size_pt=10.5)
        
        desc_lbl = p.add_run("  O que faz: ")
        apply_text_styling(desc_lbl, size_pt=10.5, bold=True)
        desc_val = p.add_run(f"{r['descricao']}\n")
        apply_text_styling(desc_val, size_pt=10.5)
        
        ben_lbl = p.add_run("  Benefício para o Comitê: ")
        apply_text_styling(ben_lbl, size_pt=10.5, bold=True, color_rgb=RGBColor(0, 128, 0))
        ben_val = p.add_run(f"{r['beneficio']}\n")
        apply_text_styling(ben_val, size_pt=10.5)
        
    # Tabela Resumida
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("\n2. Resumo de Atuação e Maturidade Tecnológica")
    apply_text_styling(h2_run, size_pt=14, bold=True, color_rgb=RGBColor(0, 32, 96))
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['Funcionalidade de IA', 'Foco do Processo', 'Impacto nos Custos / Tempo']
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_background(hdr_cells[idx], "002060")
        for p in hdr_cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                apply_text_styling(run, size_pt=10.5, bold=True, color_rgb=RGBColor(255, 255, 255))
                
    tab_data = [
        ("Triagem Inicial", "Filtro de ligações de ontem", "Altíssimo - Evita desperdício de chamadas pagas à IA"),
        ("Transcrição e Redundância", "Conversão áudio-texto", "Médio - Garante 99.8% de uptime operacional"),
        ("Pontuação de Critérios", "Nota de qualidade automatizada", "Altíssimo - Auditoria de 100% dos canais em segundos"),
        ("Parser de Chat", "Leitura de histórico de chats PDF", "Médio - Consolidação de canais falados e escritos"),
        ("Calibração / Gabarito", "Alinhamento das notas da IA", "Alto - Garante estabilidade nas regras de notas de IA")
    ]
    
    for row_data in tab_data:
        row_cells = table.add_row().cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for p in row_cells[col_idx].paragraphs:
                for run in p.runs:
                    apply_text_styling(run, size_pt=10)
                    
    doc.save(filepath)

def create_sentinel_docx(filepath):
    doc = docx.Document()
    
    # Titulo Principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("MAPA DE RECURSOS DE INTELIGÊNCIA ARTIFICIAL\nPROJETO SENTINEL")
    apply_text_styling(title_run, size_pt=16, bold=True, color_rgb=RGBColor(0, 32, 96))
    
    # Subtitulo
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Relatório Executivo de Oitivas, Análises de Imagens e Processamento de Vídeos\n")
    apply_text_styling(sub_run, size_pt=11, italic=True, color_rgb=RGBColor(100, 100, 100))
    
    # Introducao
    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run(
        "Este documento apresenta os recursos de Inteligência Artificial operados na plataforma Sentinel. "
        "O Sentinel atua como um motor de processamento de depoimentos em áudio e vídeo (oitivas) de investigações, "
        "assim como na vistoria de imagens e vídeos de danos/sinistros de cargas e veículos para a emissão padronizada de laudos."
    )
    apply_text_styling(intro_run, size_pt=11)
    
    # Secao 1: Recursos de IA
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Recursos de Inteligência Artificial e Atuação")
    apply_text_styling(h1_run, size_pt=14, bold=True, color_rgb=RGBColor(0, 32, 96))
    
    recursos = [
        {
            "nome": "Transcrição e Diarização de Oitivas / Depoimentos (Áudio)",
            "atuacao": "Upload e decodificação de depoimentos de auditoria em áudio.",
            "descricao": "IA dedicada a converter gravações longas de áudio de depoimentos em texto escrito. Realiza a identificação automática dos diferentes interlocutores (auditor, declarante, testemunha) e marca a cronologia exata das falas.",
            "beneficio": "Redução drástica do trabalho manual de escuta e digitação das declarações. Transforma horas de gravação em depoimentos estruturados e legíveis em minutos."
        },
        {
            "nome": "Geração Automática de Laudo de Depoimento (Claim-Review)",
            "atuacao": "Fechamento e consolidação técnica de investigações.",
            "descricao": "Após a homologação do texto transcrito, o motor de IA generativa (Azure OpenAI) redige um laudo técnico estruturado, resumindo contradições nos relatos, fatos confirmados, horários de ocorrências e conclusões investigativas de forma neutra.",
            "beneficio": "Padronização de laudos técnicos, velocidade na liquidação de sinistros e minimização de erros de digitação ou omissão de dados críticos da oitiva."
        },
        {
            "nome": "Visão Computacional para Análise Técnica de Imagens (Fotos)",
            "atuacao": "Módulo de Imagens e Vistorias Físicas.",
            "descricao": "IA capaz de receber fotos e imagens técnicas (fotos de avarias em veículos, danos patrimoniais, galpões e lacres de segurança). Ela analisa os detalhes visuais e descreve de forma estruturada as evidências técnicas encontradas no arquivo de imagem.",
            "beneficio": "Agilidade nas auditorias visuais de sinistro, suporte objetivo para detecção de anomalias ou fraudes em fotos enviadas e relatórios visuais gerados instantaneamente."
        },
        {
            "nome": "Análise Comportamental e Detecção de Inconsistências em Vídeo (Oitivas Gravadas)",
            "atuacao": "Módulo de Vídeo / Investigação e Oitivas de Depoimento.",
            "descricao": "IA Generativa e de Visão Computacional aplicada a depoimentos gravados em vídeo. A IA rastreia microexpressões de desconforto do declarante, analisa o distanciamento psicológico na fala (ex: uso excessivo de 'nós' em vez de 'eu' para se afastar da responsabilidade de uma ação) e detecta indícios de contradições, mentiras ou alegações inconsistentes (como falsas memórias).",
            "beneficio": "Ferramenta de alta precisão para investigadores e auditores identificarem se um depoimento possui pontos de omissão, insegurança ou incoerência comportamental de forma científica."
        },
        {
            "nome": "Análise Técnica de Vídeos Operacionais e de Trânsito (Dashcam / Câmeras de Veículos)",
            "atuacao": "Módulo de Vídeo / Sinistros, Logística e Vistorias Físicas.",
            "descricao": "IA que analisa gravações de vídeo de câmeras automotivas (dashcams) ou filmagens de cenários de acidentes e vistorias de galpões/cargas. A IA mapeia os acontecimentos cronologicamente com marcação de tempo (timestamps), descrevendo de forma minuciosa o que é observável e classificando os riscos.",
            "beneficio": "Geração automatizada de laudos de sinistros automobilísticos e vistorias físicas, convertendo minutos de vídeo em um parecer com classificação de risco (Baixo, Médio, Alto) de forma instantânea."
        }
    ]
    
    for r in recursos:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        
        name_run = p.add_run(f"• {r['nome']}\n")
        apply_text_styling(name_run, size_pt=12, bold=True, color_rgb=RGBColor(56, 189, 248))
        
        act_lbl = p.add_run("  Onde atua: ")
        apply_text_styling(act_lbl, size_pt=10.5, bold=True)
        act_val = p.add_run(f"{r['atuacao']}\n")
        apply_text_styling(act_val, size_pt=10.5)
        
        desc_lbl = p.add_run("  O que faz: ")
        apply_text_styling(desc_lbl, size_pt=10.5, bold=True)
        desc_val = p.add_run(f"{r['descricao']}\n")
        apply_text_styling(desc_val, size_pt=10.5)
        
        ben_lbl = p.add_run("  Benefício para o Comitê: ")
        apply_text_styling(ben_lbl, size_pt=10.5, bold=True, color_rgb=RGBColor(0, 128, 0))
        ben_val = p.add_run(f"{r['beneficio']}\n")
        apply_text_styling(ben_val, size_pt=10.5)
        
    # Tabela Resumida
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("\n2. Resumo de Atuação e Tecnologias Cognitivas")
    apply_text_styling(h2_run, size_pt=14, bold=True, color_rgb=RGBColor(0, 32, 96))
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['Funcionalidade de IA', 'Foco do Processo', 'Impacto nos Custos / Tempo']
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        set_cell_background(hdr_cells[idx], "002060")
        for p in hdr_cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                apply_text_styling(run, size_pt=10.5, bold=True, color_rgb=RGBColor(255, 255, 255))
                
    tab_data = [
        ("Transcrição & Diarização", "Transcrição de depoimentos e oitivas", "Reduz em 90% a digitação de depoimentos longos"),
        ("Laudo de Depoimento", "Geração de relatórios descritivos de oitivas", "Padronização e celeridade nos pareceres técnicos"),
        ("Visão Computacional (Fotos)", "Análise de fotos de avarias e cargas", "Apoio visual rápido para atestar danos físicos e sinistros"),
        ("Análise de Vídeo Comportamental", "Oitivas e depoimentos gravados em vídeo", "Identificação de contradições e microexpressões corporais"),
        ("Análise de Vídeo de Dashcam", "Vistorias e câmeras automotivas", "Laudo cronológico de acidentes e classificação de risco")
    ]
    
    for row_data in tab_data:
        row_cells = table.add_row().cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for p in row_cells[col_idx].paragraphs:
                for run in p.runs:
                    apply_text_styling(run, size_pt=10)
                    
    doc.save(filepath)

def main():
    desktop_dir = os.path.join(os.path.expanduser("~"), "Desktop")
    base_dir = os.path.join(desktop_dir, "recursos-ia")
    
    auditoria_dir = os.path.join(base_dir, "auditoria")
    sentinel_dir = os.path.join(base_dir, "sentinel")
    
    os.makedirs(auditoria_dir, exist_ok=True)
    os.makedirs(sentinel_dir, exist_ok=True)
    
    print(f"Diretorios criados no Desktop sob: {base_dir}")
    
    auditoria_docx = os.path.join(auditoria_dir, "Recursos_de_IA_Sistema_Auditoria.docx")
    sentinel_docx = os.path.join(sentinel_dir, "Recursos_de_IA_Projeto_Sentinel.docx")
    
    create_auditoria_docx(auditoria_docx)
    print(f"Criado com sucesso: {auditoria_docx}")
    
    create_sentinel_docx(sentinel_docx)
    print(f"Criado com sucesso: {sentinel_docx}")

if __name__ == "__main__":
    main()
