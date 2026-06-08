import os
import subprocess
import sys

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

ensure_docx()

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_report():
    doc = Document()
    
    title = doc.add_heading('Relatório Técnico: Extração de Pesos, Monitoramento Azure e Histórico de Sincronização', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('Data: 27 de Abril de 2026').bold = True
    doc.add_paragraph('Projeto: Auditoria').bold = True
    doc.add_paragraph('Diretório Base: C:\\Users\\lucas.afonso\\projetos\\auditoria').bold = True
    
    doc.add_paragraph('---')
    
    # 1. Contexto e Objetivos
    doc.add_heading('1. Contexto e Objetivos', level=1)
    p = doc.add_paragraph('Dando sequência às melhorias no módulo de Telefonia e Qualidade da Auditoria, este ciclo focou em três pilares principais para preparar o sistema para a operação assistida por IA (GPT-4o) em larga escala:')
    doc.add_paragraph('1. Estruturação de Critérios: Refatorar a extração de pesos e regras do Excel para um formato JSON universal.')
    doc.add_paragraph('2. Monitoramento Cognitivo: Avaliar a saúde e a conectividade dos serviços de IA na Azure.')
    doc.add_paragraph('3. Observabilidade da Telefonia: Substituir os dados em memória ("placeholder") da API de Sincronização por uma tabela real no banco de dados.')
    
    # 2. Atividade 1
    doc.add_heading('2. Atividade 1: Script de Extração de Pesos (JSON)', level=1)
    doc.add_heading('2.1. Problema Original', level=2)
    doc.add_paragraph('O script extract_pesos_detailed.py era um simples explorador em console (usando pandas). Ele imprimia no terminal os critérios por setor, mas não gerava nenhum artefato que o backend pudesse consumir. Isso obrigava o backend a utilizar regras ou pesos hardcoded.')
    
    doc.add_heading('2.2. Solução Implementada', level=2)
    doc.add_paragraph('- O script foi completamente reescrito para ler a planilha CRITÉRIOS - PESOS -.xlsm.', style='List Bullet')
    doc.add_paragraph('- Foi adicionada lógica para sanitização de dados: tratamento de valores numéricos nulos (NaN para 0.0) e extração baseada em Expressões Regulares.', style='List Bullet')
    doc.add_paragraph('- Exportação: O resultado agora é gravado automaticamente no arquivo criterios_pesos_extraidos.json contendo um dicionário de categorias (Mondelez, Logística, Cadastro, etc.), onde cada categoria possui uma lista de objetos contendo referencia, pergunta, peso e deflator.', style='List Bullet')
    doc.add_paragraph('- Status: Script rodou com sucesso extraindo 38 categorias.', style='List Bullet')
    
    # 3. Atividade 2
    doc.add_heading('3. Atividade 2: Monitoramento Azure e Avaliação Cognitiva', level=1)
    doc.add_heading('3.1. Execução do check_azure_health.py', level=2)
    doc.add_paragraph('Para garantir que a integração da telefonia seria processada corretamente, rodamos o verificador de APIs cognitivas da Microsoft. Os resultados foram:')
    doc.add_paragraph('- GPT-4o Principal (gpt-4o): ✅ Conectado com sucesso! (Vital para o julgamento da auditoria).', style='List Bullet')
    doc.add_paragraph('- Speech Services (eastus2): ✅ Autorizado com sucesso! (Vital para processamento inicial de áudio).', style='List Bullet')
    doc.add_paragraph('- Whisper Legado (nstech-bas-whisper): ❌ FALHA 401 - Access denied. A chave de assinatura antiga da instância de Whisper expirou ou foi rotacionada. Obs: Como o sistema principal de produção já havia migrado para usar gpt-4o-transcribe-diarize na Foundry (que é mais seguro e faz diarização), esta falha no Whisper não quebra a operação atual, mas alerta para a necessidade de limpar as chaves mortas do .env.', style='List Bullet')
    doc.add_paragraph('- Text Analytics: Chaves ausentes no .env. Recurso não essencial neste momento.', style='List Bullet')
    
    # 4. Atividade 3
    doc.add_heading('4. Atividade 3: Banco de Dados de Histórico de Telefonia', level=1)
    doc.add_heading('4.1. Problema Original', level=2)
    doc.add_paragraph('O endpoint da API /api/telefonia/sync/history estava retornando apenas a última sincronização que ficava salva em uma variável de memória RAM (_LAST_SYNC). Quando o Cloud Run era reiniciado, o histórico era completamente perdido.')
    
    doc.add_heading('4.2. Modelagem e Banco de Dados (PostgreSQL)', level=2)
    doc.add_paragraph('- Migração Criada: m20260427_001_create_huawei_sync_history.', style='List Bullet')
    doc.add_paragraph('- Tabela: telefonia_sync_history.', style='List Bullet')
    doc.add_paragraph('- Colunas criadas: id, started_at, finished_at, status, horas_retroativas, baixadas, enfileiradas, erros_totais, mensagem_erro e trigger_type (identificando se rodou por cron ou manual).', style='List Bullet')
    doc.add_paragraph('- O script db_migrate.py foi executado para atualizar a base local.', style='List Bullet')
    
    doc.add_heading('4.3. Refatoração do Backend', level=2)
    doc.add_paragraph('- Repositório Adicionado: Criamos backend/repositories/telefonia.py contendo funções nativas (save_telefonia_sync_history, list_telefonia_sync_history).', style='List Bullet')
    doc.add_paragraph('- Proxy (database.py): As funções de salvamento e listagem foram injetadas para que toda a aplicação possa usar sem acoplamento direto de repositório.', style='List Bullet')
    doc.add_paragraph('- Roteadores (telefonia.py): Atualizamos as funções assíncronas _run_manual_sync e cron_sync para invocar a escrita no banco no final de cada lote de extração (sucesso ou falha). A API /sync/history foi limpa do seu estado "placeholder" e passou a extrair os dados diretamente do novo histórico em PostgreSQL.', style='List Bullet')
    
    doc.add_heading('4.4. Estabilidade', level=2)
    doc.add_paragraph('Os testes unitários foram re-executados garantindo preservação de escopo. As melhorias na telefonia e qualidade não geraram impactos nas camadas sensíveis de revisão e usuários.')
    
    # 5. Atualização de Skills
    doc.add_heading('5. Atualização de Skills (Agentes)', level=1)
    doc.add_paragraph('O skill auditoria-management foi atualizado para instruir agentes futuros a estarem cientes do uso de criterios_pesos_extraidos.json para regras do Negócio, alertá-los de que o Whisper Legado está em deprecation (a favor do GPT-4o-transcribe), e formalizar a tabela telefonia_sync_history como fonte da verdade de auditoria de consumo da Huawei.')
    
    os.makedirs('relatorios_atividades', exist_ok=True)
    file_path = os.path.join('relatorios_atividades', 'melhorias_apis_e_criterios.docx')
    doc.save(file_path)
    print(f"Documento Word gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    create_word_report()
