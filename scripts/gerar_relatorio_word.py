import os
import subprocess
import sys

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

ensure_docx()

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_report():
    doc = Document()
    
    # Título principal
    title = doc.add_heading('Relatório Técnico: Integração Telefonia Huawei e Resolução de Anomalias', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadados
    doc.add_paragraph('Data: 27 de Abril de 2026').bold = True
    doc.add_paragraph('Projeto: Auditoria').bold = True
    doc.add_paragraph('Diretório Base: C:\\Users\\lucas.afonso\\projetos\\auditoria').bold = True
    
    doc.add_paragraph('---')
    
    # 1. Contexto e Objetivos
    doc.add_heading('1. Contexto e Objetivos', level=1)
    doc.add_paragraph('O objetivo desta frente de trabalho foi estabilizar, diagnosticar e corrigir as integrações de rede com a API da Huawei (AICC). A prioridade era resolver o bloqueio de rede (Erro HTTP 403) reportado pelo serviço do Google Cloud Run e corrigir a anomalia de paginação que ocultava as gravações sincronizadas da interface web de telefonia.')
    
    # 2. Atividade 1
    doc.add_heading('2. Atividade 1: Investigação e Resolução do Erro 403 (Autenticação Huawei)', level=1)
    
    doc.add_heading('2.1. Diagnóstico Inicial', level=2)
    doc.add_paragraph('- Sintoma: O arquivo de logs full_logs.json e as execuções do Cron registravam erros 403 Forbidden oriundos da comunicação com a Huawei.', style='List Bullet')
    doc.add_paragraph('- Topologia de Rede: O sistema estava operando sob HUAWEI_AUTH_MODE=proxy, onde o tráfego era roteado e assinado via um endpoint intermediário na Teledata (c2Authorization.php).', style='List Bullet')
    doc.add_paragraph('- Ação Desejada: O provedor da Huawei reportou a inclusão do IP estático do nosso Google Cloud NAT (35.199.111.152) em sua whitelist para acesso direto (usando fluxo OAuth 2.0 e o endpoint tokenByAkSk). O objetivo era testar a viabilidade do acesso direto sem quebrar a integração de produção atual.', style='List Bullet')
    
    doc.add_heading('2.2. Implementação do Teste de Prova de Conceito (PoC)', level=2)
    doc.add_paragraph('- Desenvolvemos um script isolado backend/tmp_test_huawei_oauth.py implementando estritamente a mesma cadeia HTTP da coleção do Postman (Huawei AICC BrazilSaaS-OPENTECH Lucas.postman_collection.json).', style='List Bullet')
    doc.add_paragraph('- Fizemos o deploy do script na nuvem via Cloud Run Job (test-huawei-oauth), com o parâmetro --vpc-egress=all-traffic, garantindo que o teste utilizaríamos obrigatoriamente o IP 35.199.111.152 para atingir o WAF da Huawei.', style='List Bullet')
    
    doc.add_heading('2.3. Execuções e Resultados (Logs)', level=2)
    p1 = doc.add_paragraph('- Tentativa 1 (Falha 403): Ao invocar a API com as variáveis de ambiente atuais do sistema (HUAWEI_AK: 97501...), recebemos novamente o bloqueio.', style='List Bullet')
    doc.add_paragraph('  - Log Obtido: Status Auth: 403 - Falha na Autenticação.')
    doc.add_paragraph('  - Hipótese Inicial: Bloqueio de WAF por falta de propagação da whitelist do IP.')
    
    doc.add_paragraph('- Ação Corretiva (Engenharia): Em testes paralelos e validação da coleção do Postman, o desenvolvedor (Lucas) notou que o app_key necessário para o endpoint de autenticação direta era estruturalmente diferente (formato UUID) do app_key esperado pelo proxy da Teledata.', style='List Bullet')
    
    doc.add_paragraph('- Tentativa 2 (Sucesso 200): Alteramos o script para fazer hardcode das chaves diretas do Postman (app_key: <HUAWEI_DIRECT_APP_KEY>...).', style='List Bullet')
    doc.add_paragraph('  - Log Obtido: Status Auth: 200 - Token recebido com sucesso!')
    doc.add_paragraph('  - Veredito: A rede e o IP 35.199.111.152 estão confirmadamente liberados no ambiente da Huawei. O falso-positivo de bloqueio ocorria porque enviávamos credenciais desconhecidas (da Teledata) para a porta principal da Huawei, resultando em rejeição pelo Application Gateway.')
    
    doc.add_heading('2.4. Próximos Passos Obrigatórios (Implementação Futura)', level=2)
    doc.add_paragraph('- Variáveis de Ambiente: Adicionar novas credenciais exclusivas no .env do Cloud Run: HUAWEI_DIRECT_APP_KEY e HUAWEI_DIRECT_APP_SECRET.', style='List Bullet')
    doc.add_paragraph('- Refatoração no Backend: O arquivo backend/core/huawei_client.py deve ser modificado para suportar o modo oauth_direct, onde o AccessToken retornado de tokenByAkSk é utilizado no Header Authorization: Bearer <token>. Isso permitirá desativar definitivamente a dependência da URL de proxy da Teledata.', style='List Bullet')
    
    doc.add_paragraph('---')
    
    # 3. Atividade 2
    doc.add_heading('3. Atividade 2: Resolução da Anomalia de Paginação (Fila de Triagem)', level=1)
    
    doc.add_heading('3.1. Diagnóstico do Problema Visual', level=2)
    doc.add_paragraph('- Sintoma Relatado: As ligações telefônicas eram reportadas como baixadas com sucesso e enfileiradas nos logs da execução (_enfileirar_classificado), mas a tabela na interface do Cloud Run (/api/telefonia/recordings) era carregada vazia.', style='List Bullet')
    
    doc.add_heading('3.2. Causa Raiz Técnica (Anti-padrão de Filtragem)', level=2)
    doc.add_paragraph('- A rota no backend (routers/telefonia.py) solicitava os dados executando database.listar_fila_revisao_classificacao(limit=50).', style='List Bullet')
    doc.add_paragraph('- O PostgreSQL resolvia essa requisição trazendo as últimas 50 ligações enfileiradas de qualquer origem (o que incluía massa de dados gerada por uploads manuais em PDF/Excel).', style='List Bullet')
    doc.add_paragraph('- Em seguida, o Python rodava uma List Comprehension para filtrar apenas aquelas originadas pela telefonia: if origem == \'huawei_sync\'.', style='List Bullet')
    doc.add_paragraph('- O Bug: Como os 50 registros recortados do banco não possuíam gravações da Huawei (pois estas estavam mais para trás na fila, nas posições 51+), o Python renderizava uma lista vazia, simulando falsamente a ausência dos áudios no sistema.', style='List Bullet')
    
    doc.add_heading('3.3. Solução Implementada', level=2)
    doc.add_paragraph('A lógica de filtragem foi movida inteiramente da aplicação (memória) para a camada de Banco de Dados (PostgreSQL):')
    doc.add_paragraph('1. Camada de Repositório (repositories/classification_review.py):')
    doc.add_paragraph('   - Injetamos o parâmetro opcional origem.')
    doc.add_paragraph('   - Adicionada condição nativa de consulta em campo JSONB: filtros.append("metadata_json::jsonb ->> \'origem\' = %s").')
    doc.add_paragraph('2. Camada de Rotas (routers/telefonia.py):')
    doc.add_paragraph('   - A List Comprehension ineficiente foi removida.')
    doc.add_paragraph('   - O argumento origem=\'huawei_sync\' passou a ser repassado ao DAO, forçando o banco de dados a descartar as ligações manuais antes de aplicar o LIMIT 50.')
    
    doc.add_heading('3.4. Testes, Tratamento de Conflitos e Lançamento', level=2)
    doc.add_paragraph('- Testes de Contrato: Um erro transitório de import (REVIEW_QUEUE_STATUS_PENDING) em database.py foi exposto pelos testes unitários e prontamente corrigido. Todos os 12 testes do Módulo de Fila finalizaram em PASSED.', style='List Bullet')
    doc.add_paragraph('- Resolução Git: Um Merge Conflict severo (HEAD vs Origin) ocorreu em routers/telefonia.py durante o git pull --rebase. O conflito foi resolvido cirurgicamente priorizando a nova injeção JSONB.', style='List Bullet')
    doc.add_paragraph('- Deploy e Sucesso:', style='List Bullet')
    doc.add_paragraph('  - Código commitado e mesclado: 7f32301 fix(telefonia): corrige filtro de origens no banco de dados para evitar paginacao vazia na fila de revisao.')
    doc.add_paragraph('  - Lançamento no Cloud Run efetuado via gcloud run deploy auditoria --source .')
    doc.add_paragraph('  - A revisão ativa provê agora exibição fluida e contínua das auditorias importadas via Huawei, ignorando poluição de dados manuais na fila.')
    
    doc.add_paragraph('---')
    
    # 4. Orientações Gerais
    doc.add_heading('4. Orientações Gerais para Próximos Agentes de IA', level=1)
    doc.add_paragraph('1. Autenticação: Não submeter payloads com credenciais da Teledata (HUAWEI_AK: 975...) diretamente à brazilsaas.aicccloud.com. Use apenas o UUID associado ao endpoint de governança.', style='List Number')
    doc.add_paragraph('2. Manipulação de Filas (Performance): É expressamente proibido usar List Comprehensions do Python para filtrar dados da Fila de Revisão que já foram paginados/limitados. Todo filtro deve ser injetado no Postgres.', style='List Number')
    doc.add_paragraph('3. Ambiente de Teste: Caso seja necessário rodar diagnósticos restritos de firewall contra parceiros (ex: Huawei), instancie sempre um Job no Cloud Run amarrado à auditoria-vpc-connector para garantir saída com IP Estático mapeado.', style='List Number')
    
    # Salvar o documento
    os.makedirs('relatorios_atividades', exist_ok=True)
    file_path = os.path.join('relatorios_atividades', 'telefonia_huawei.docx')
    doc.save(file_path)
    print(f"Documento Word gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    create_word_report()
