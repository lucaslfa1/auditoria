import os
import subprocess
import sys

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("Instalando python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

ensure_docx()

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_word_report():
    doc = Document()
    
    title = doc.add_heading('Relatório de Atualização: Fixação de Busca e Sincronização Huawei AICC', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('Data: 27 de Abril de 2026').bold = True
    doc.add_paragraph('Projeto: Auditoria').bold = True
    doc.add_paragraph('Autor da Revisão: Codex GPT 5.5 / Gemini CLI').bold = True
    
    doc.add_paragraph('---')
    
    doc.add_heading('1. Visão Executiva', level=1)
    doc.add_paragraph('Este documento registra a resolução dos gargalos críticos finais que impediam o download efetivo dos áudios da Huawei AICC para a nossa fila de revisão de telefonia. Com a aplicação destas correções, o sistema de integração foi ativado com sucesso em ambiente real, alcançando o marco de 200 chamadas únicas na VDN em um intervalo de 1 hora.')
    
    doc.add_heading('2. Auditoria e Correções no Banco de Dados', level=1)
    doc.add_paragraph('Problema: A aplicação falhava no momento de registrar o log final da sincronização devido à divergência de colunas.')
    doc.add_paragraph('- Causa Raiz: A tabela huawei_sync_logs não continha as colunas status e failure_reason.', style='List Bullet')
    doc.add_paragraph('- Ação Aplicada: Foi criada e executada uma nova migração de banco de dados (m20260427_002_add_status_to_huawei_sync_logs), ajustando a estrutura. O erro de registro de sincronização foi totalmente extinto.', style='List Bullet')
    
    doc.add_heading('3. Auditoria na Regra de Busca de Janela Temporal (Bug 0100002)', level=1)
    doc.add_paragraph('Problema: Chamadas da Huawei rejeitadas pelo código de erro 0100002 (Tempo Inválido).')
    doc.add_paragraph('- Causa Raiz: A lógica de cálculo do limite temporal da janela de busca gerava uma sobra de 1 milissegundo, o qual a API do AICC considerava um formato de data/hora malformado.', style='List Bullet')
    doc.add_paragraph('- Ação Aplicada: Aritmética do recorte de janela foi corrigida no código de sincronização (huawei_sync.py).', style='List Bullet')
    doc.add_paragraph('- Validação de Leitura: A consulta re-testada retornou precisamente 200 chamadas validadas em uma janela temporal consistente (200 callIds únicos).', style='List Bullet')
    
    doc.add_heading('4. Teste de Sincronização Real (Status Atual)', level=1)
    doc.add_paragraph('Os testes realizados no ambiente de produção após os patches indicaram a seguinte performance de funil em um ciclo de 1 hora:')
    doc.add_paragraph('- Total na VDN: 200 chamadas.', style='List Bullet')
    doc.add_paragraph('- Válidas pós-filtro: 20 chamadas (Apenas os perfis alvo que a auditoria deve ouvir).', style='List Bullet')
    doc.add_paragraph('- Candidatos a Download: 20.', style='List Bullet')
    doc.add_paragraph('- Áudios Baixados: 4.', style='List Bullet')
    doc.add_paragraph('- Enfileiradas na IA: 2.', style='List Bullet')
    doc.add_paragraph('- Descartadas por Duplicação: 2.', style='List Bullet')
    doc.add_paragraph('- Erros de Sync: 0 ([]).', style='List Bullet')
    
    doc.add_heading('5. Mapeamento de Gargalos Restantes (Próximos Passos Obrigatórios)', level=1)
    doc.add_paragraph('Apesar do fluxo de rede e roteamento estarem perfeitos (conforme funil na Seção 4), a extração caiu bruscamente de "20 candidatos" para apenas "4 baixados".')
    doc.add_paragraph('- Motivo Diagnosticado: A API principal da Huawei não encontrou o dado bruto para a maioria dos áudios e retornou o erro 0300012 No data found no momento do download.', style='List Bullet')
    doc.add_paragraph('- Ação Requerida (Fallback OBS): A infraestrutura possui um mecanismo de contorno (já codificado) baseado na leitura do storage nativo da Huawei (OBS). Porém, a solução de Fallback não foi engatilhada em produção porque as variáveis ambientais do OBS não estão declaradas no sistema.', style='List Bullet')
    doc.add_paragraph('- Recomendação: Solicitar imediatamente à Operação/Teledata as credenciais do bucket OBS da telefonia e alimentá-las na configuração do Cloud Run para assegurar o download de 100% das ligações retidas: HUAWEI_OBS_AK, HUAWEI_OBS_SK, HUAWEI_OBS_BUCKET, HUAWEI_OBS_ENDPOINT.', style='List Bullet')
    
    doc.add_heading('6. Certificação e Baseline', level=1)
    doc.add_paragraph('Após as alterações promovidas pelo agente, todo o módulo de Quality Assurance foi testado positivamente sem quebras de contratos de software:')
    doc.add_paragraph('- Suíte Integrada (pytest): Os testes em test_huawei_client.py, test_huawei_sync.py e test_review_queue_contract.py reportaram 39 PASSED.', style='List Bullet')
    doc.add_paragraph('- Compilação PyCache local dos arquivos alterados: OK.', style='List Bullet')
    
    os.makedirs('relatorios_atividades', exist_ok=True)
    file_path = os.path.join('relatorios_atividades', 'atualizacao_codex_huawei_sync.docx')
    doc.save(file_path)
    print(f"Documento Word gerado com sucesso em: {file_path}")

if __name__ == "__main__":
    create_word_report()
