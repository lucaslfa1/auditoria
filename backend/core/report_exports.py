"""Geração de relatórios exportáveis de um resultado de auditoria.

Papel no sistema: converte um `AuditResult` (resultado de auditoria já avaliado)
em arquivos baixáveis nos formatos Excel (.xlsx), Word (.docx) e PDF — tanto do
relatório de critérios/notas quanto da transcrição da ligação. Cada função
retorna um buffer em memória (`io.BytesIO`) já posicionado no início, pronto
para ser servido como download pela camada de API.

As bibliotecas pesadas (pandas/openpyxl, python-docx, reportlab) são importadas
preguiçosamente dentro de cada função para não onerar o boot do processo.

Sem custo de API: só formatação/CPU em memória. Não chama Azure OpenAI/Speech,
banco nem rede.
"""
from __future__ import annotations

from datetime import datetime
import io
from xml.sax.saxutils import escape

from schemas import AuditResult


# Rótulos exibidos nos relatórios para o status de cada critério (contrato de UI
# em PT-BR). Apenas dois estados visíveis: "Atende" / "Não atende".
STATUS_LABELS = {
    'pass': 'Atende',
    'fail': 'Não atende',
}


def _format_detail_status(status: str) -> str:
    """Mapeia o status técnico de um critério para o rótulo binário exibido.

    Normaliza o valor cru (pass/na/n/a/pending_manual -> "Atende";
    fail/partial e qualquer outro -> "Não atende"). Note que `partial` é
    apresentado como reprovação no relatório.
    """
    normalized = str(status or '').strip().lower()
    if normalized in {'pass', 'na', 'n/a', 'pending_manual'}:
        return STATUS_LABELS['pass']
    if normalized in {'fail', 'partial'}:
        return STATUS_LABELS['fail']
    return STATUS_LABELS['fail']


def _pdf_text(value: object) -> str:
    """Escapa um valor para uso seguro em parágrafos do reportlab.

    Aplica escape de XML/HTML e converte quebras de linha em ``<br/>`` (markup
    aceito pelo `Paragraph` do reportlab).
    """
    return escape(str(value or '')).replace('\n', '<br/>')


def generate_excel_report(result: AuditResult):
    """Gera o relatório de auditoria em Excel (.xlsx).

    Produz uma planilha com duas abas: "Resumo" (data, operador, nota final e
    resumo) e "Detalhes" (uma linha por critério: rótulo, status, peso, nota
    obtida e comentário), com larguras de coluna e quebra de texto ajustadas.

    Retorna um `io.BytesIO` posicionado no início. Sem efeitos colaterais
    externos (escreve só em memória).
    """
    import pandas as pd
    from openpyxl.styles import Alignment

    df_header = pd.DataFrame(
        {
            'Data': [datetime.now().strftime('%d/%m/%Y %H:%M')],
            'Operador': [result.operatorName],
            'Nota Final': [f'{result.score} / {result.maxPossibleScore}'],
            'Resumo': [result.summary],
        }
    )
    df_details = pd.DataFrame(
        [
            {
                'Critério': detail.label,
                'Status': _format_detail_status(detail.status),
                'Peso': detail.weight,
                'Nota Obtida': detail.obtainedScore,
                'Comentário': detail.comment,
            }
            for detail in result.details
        ]
    )
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_header.to_excel(writer, sheet_name='Resumo', index=False)
        df_details.to_excel(writer, sheet_name='Detalhes', index=False)

        ws_resumo = writer.sheets['Resumo']
        ws_resumo.column_dimensions['A'].width = 18
        ws_resumo.column_dimensions['B'].width = 30
        ws_resumo.column_dimensions['C'].width = 15
        ws_resumo.column_dimensions['D'].width = 80

        ws_detalhes = writer.sheets['Detalhes']
        ws_detalhes.column_dimensions['A'].width = 50
        ws_detalhes.column_dimensions['B'].width = 16
        ws_detalhes.column_dimensions['C'].width = 10
        ws_detalhes.column_dimensions['D'].width = 15
        ws_detalhes.column_dimensions['E'].width = 80

        default_align = Alignment(wrap_text=True, vertical='center')
        for sheet_name in ['Resumo', 'Detalhes']:
            worksheet = writer.sheets[sheet_name]
            for row in worksheet.iter_rows():
                for cell in row:
                    cell.alignment = default_align
    output.seek(0)
    return output


def generate_docx_report(result: AuditResult):
    """Gera o relatório de auditoria em Word (.docx).

    Monta um documento com cabeçalho (data, operador, nota final, resumo) e uma
    tabela de critérios (Critério, Status, Peso, Nota Obtida, Comentário).

    Retorna um `io.BytesIO` posicionado no início. Escreve só em memória.
    """
    from docx import Document

    doc = Document()
    doc.add_heading('Relatório de Auditoria', level=1)
    doc.add_paragraph(
        f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
        f'Operador: {result.operatorName}\n'
        f'Nota Final: {result.score} / {result.maxPossibleScore}\n'
        f'Resumo: {result.summary or ""}'
    )
    table = doc.add_table(rows=1, cols=5)
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text, header[3].text, header[4].text = (
        'Critério',
        'Status',
        'Peso',
        'Nota Obtida',
        'Comentário',
    )
    for detail in result.details:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = (
            detail.label,
            _format_detail_status(detail.status),
            str(detail.weight),
            str(detail.obtainedScore),
            detail.comment,
        )
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_docx_transcription(result: AuditResult):
    """Gera a transcrição da ligação em Word (.docx).

    Escreve um parágrafo por segmento de `result.transcription`, no formato
    "{start} - {end} | {text}".

    Retorna um `io.BytesIO` posicionado no início. Escreve só em memória.
    """
    from docx import Document

    doc = Document()
    doc.add_heading('Transcricao da Ligacao', level=1)
    for segment in result.transcription:
        doc.add_paragraph(f'{segment.start} - {segment.end} | {segment.text}')
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_pdf_report(result: AuditResult):
    """Gera o relatório de auditoria em PDF (A4, via reportlab).

    Renderiza cabeçalho (data, operador, nota), seção de resumo e uma tabela de
    critérios (Critério, Status, Nota Obtida, Comentário) com cabeçalho repetido
    a cada página. Textos passam por `_pdf_text` (escape + quebras de linha).

    Retorna um `io.BytesIO` posicionado no início. Escreve só em memória.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ExportBody',
        parent=styles['BodyText'],
        fontSize=9,
        leading=12,
        wordWrap='CJK',
    ))
    styles.add(ParagraphStyle(
        name='ExportSmall',
        parent=styles['BodyText'],
        fontSize=8,
        leading=10,
        wordWrap='CJK',
    ))
    styles.add(ParagraphStyle(
        name='ExportTitle',
        parent=styles['Title'],
        fontSize=16,
        leading=20,
        wordWrap='CJK',
    ))

    story = [
        Paragraph('Relatório de Auditoria', styles['ExportTitle']),
        Spacer(1, 8),
        Paragraph(f'<b>Data:</b> {_pdf_text(datetime.now().strftime("%d/%m/%Y %H:%M"))}', styles['ExportBody']),
        Paragraph(f'<b>Operador:</b> {_pdf_text(result.operatorName)}', styles['ExportBody']),
        Paragraph(f'<b>Nota:</b> {_pdf_text(result.score)} / {_pdf_text(result.maxPossibleScore)}', styles['ExportBody']),
        Spacer(1, 10),
        Paragraph('<b>Resumo</b>', styles['Heading2']),
        Paragraph(_pdf_text(result.summary), styles['ExportBody']),
        Spacer(1, 10),
    ]

    table_data = [[
        Paragraph('<b>Critério</b>', styles['ExportSmall']),
        Paragraph('<b>Status</b>', styles['ExportSmall']),
        Paragraph('<b>Nota Obtida</b>', styles['ExportSmall']),
        Paragraph('<b>Comentário</b>', styles['ExportSmall']),
    ]]
    for detail in result.details:
        table_data.append([
            Paragraph(_pdf_text(detail.label), styles['ExportSmall']),
            Paragraph(_pdf_text(_format_detail_status(detail.status)), styles['ExportSmall']),
            Paragraph(_pdf_text(detail.obtainedScore), styles['ExportSmall']),
            Paragraph(_pdf_text(detail.comment), styles['ExportSmall']),
        ])

    details_table = Table(
        table_data,
        colWidths=[4.1 * cm, 2.2 * cm, 2.2 * cm, 8.4 * cm],
        repeatRows=1,
    )
    details_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E2E8F0')),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CBD5E1')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(Paragraph('<b>Detalhes</b>', styles['Heading2']))
    story.append(details_table)

    doc.build(story)
    output.seek(0)
    return output


def generate_pdf_transcription(result: AuditResult):
    """Gera a transcrição da ligação em PDF (A4, via reportlab).

    Para cada segmento de `result.transcription`, renderiza um carimbo de tempo
    ("{start} - {end}", ou só `start` quando não há `end`) seguido do texto.

    Retorna um `io.BytesIO` posicionado no início. Escreve só em memória.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='TranscriptionTimestamp',
        parent=styles['BodyText'],
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#64748B'),
        wordWrap='CJK',
    ))
    styles.add(ParagraphStyle(
        name='TranscriptionBody',
        parent=styles['BodyText'],
        fontSize=9,
        leading=12,
        spaceAfter=6,
        wordWrap='CJK',
    ))

    story = [
        Paragraph('Transcrição da Ligação', styles['Title']),
        Spacer(1, 10),
    ]
    for segment in result.transcription:
        timestamp = f'{segment.start} - {segment.end}' if segment.end else segment.start
        story.append(Paragraph(_pdf_text(timestamp), styles['TranscriptionTimestamp']))
        story.append(Paragraph(_pdf_text(segment.text), styles['TranscriptionBody']))

    doc.build(story)
    output.seek(0)
    return output
