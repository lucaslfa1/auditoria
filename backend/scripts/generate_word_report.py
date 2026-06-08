import os
from docx import Document
from docx.shared import Pt, Inches

def create_report(output_path):
    doc = Document()

    # Título
    title = doc.add_heading('Relatório: Módulo de Ligação (Sincronização Huawei)', 0)
    title.alignment = 1

    doc.add_paragraph("Este relatório detalha a arquitetura, os arquivos envolvidos e o estado atual de funcionamento do módulo de ligação e sincronia com a telefonia Huawei no sistema de Auditoria nstech.")

    # Seção 1: Arquivos do Módulo
    doc.add_heading('1. Arquivos que compõem o módulo', level=1)
    
    doc.add_heading('1.1 Backend (Python)', level=2)
    doc.add_paragraph('Core & Integração:', style='List Bullet')
    doc.add_paragraph('backend/core/huawei_sync.py: O "cérebro" da sincronização. Orquestra a busca, filtragem (recordId) e a rotina de download.', style='List Bullet 2')
    doc.add_paragraph('backend/core/huawei_client.py: Cliente HTTP que fala com a API da Huawei (via proxy ou oauth_direct).', style='List Bullet 2')
    doc.add_paragraph('backend/core/huawei_obs_client.py: Cliente para acessar o repositório secundário (OBS/Bucket) quando o sistema primário (CC-FS) não possui a gravação.', style='List Bullet 2')
    
    doc.add_paragraph('Rotas & API:', style='List Bullet')
    doc.add_paragraph('backend/routers/telefonia.py: Expõe os endpoints para o Frontend (/api/telefonia/recordings, /api/telefonia/sync/status).', style='List Bullet 2')
    
    doc.add_paragraph('Banco de Dados & Scripts:', style='List Bullet')
    doc.add_paragraph('backend/repositories/telefonia.py: Centraliza algumas buscas e filtros de banco.', style='List Bullet 2')
    doc.add_paragraph('backend/scripts/run_huawei_sync.py: Script que pode ser rodado manualmente via terminal para engatilhar um ciclo de sincronia isolado.', style='List Bullet 2')
    
    doc.add_heading('1.2 Frontend (TypeScript/React)', level=2)
    doc.add_paragraph('src/features/telefonia/components/TelefoniaPage.tsx: A tela principal (painel) onde os administradores visualizam as ligações baixadas e pendentes.', style='List Bullet')
    doc.add_paragraph('src/features/telefonia/hooks/useTelefoniaSync.ts: O hook que controla o botão de "Iniciar Sync", gerenciando o tempo escolhido e o loading da tela.', style='List Bullet')
    doc.add_paragraph('src/features/telefonia/components/HuaweiCredentialsCard.tsx: Componente da tela de configurações para gerenciar senhas e chaves da Huawei.', style='List Bullet')

    # Seção 2: Como a Função Opera Atualmente
    doc.add_heading('2. Estado Atual e Funcionamento da Função', level=1)
    
    doc.add_heading('2.1 O que é o "Tempo Escolhido" (Janela Retroativa)', level=2)
    p = doc.add_paragraph()
    p.add_run("Quando o usuário inicia o sync pela interface escolhendo um período (ex: 1 hora), o sistema ")
    p.add_run("não").bold = True
    p.add_run(" fica ligado durante uma hora baixando as chamadas que caem. Em vez disso, ele define uma ")
    p.add_run("janela retroativa de busca").bold = True
    p.add_run(". O sistema olha para o relógio agora, subtrai 1 hora, e pede à Huawei a lista de todas as ligações que ocorreram naquele intervalo.")

    doc.add_heading('2.2 O Filtro Inteligente (A correção do "recordId")', level=2)
    p2 = doc.add_paragraph()
    p2.add_run("A lista retornada pela Huawei contém todo tipo de evento (quedas na URA, transferências, chamadas não atendidas). O commit recente de correção introduziu uma regra crucial no ")
    p2.add_run("huawei_sync.py").bold = True
    p2.add_run(": a ")
    p2.add_run("priorização por recordId").bold = True
    p2.add_run(". O recordId é a \"identidade\" do arquivo de áudio. Se ele estiver vazio, a ligação não tem áudio. O sistema agora coloca no topo da fila de download apenas as ligações com recordId válido, não desperdiçando tempo de processamento com os eventos vazios.")

    doc.add_heading('2.3 O Fluxo de Extração e Inteligência Artificial', level=2)
    doc.add_paragraph('1. Download: O sistema baixa o arquivo .wav via CC-FS ou OBS (bucket fallback).', style='List Number')
    doc.add_paragraph('2. Transcrição: O áudio é enviado para a infraestrutura Microsoft Azure (modelo GPT-4o-Audio-Preview para Diarização, com fallback para Whisper) que transforma as falas em texto separado por locutor.', style='List Number')
    doc.add_paragraph('3. Classificação: O texto é repassado ao Azure GPT-4o, que lê a conversa completa, compara com os critérios de cada setor e aplica a nota Pass/Fail.', style='List Number')

    doc.add_heading('2.4 Configurações de Rede (Auth Mode)', level=2)
    doc.add_paragraph('Atualmente, a comunicação de rede exige IP liberado na whitelist da Teledata. O sistema opera no modo "proxy", usando a URL da Teledata para validar o acesso à AICC da Huawei, usando o IP de saída seguro da Nuvem do Google (35.199.111.152).')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph("Relatório gerado automaticamente pelo assistente de Inteligência Artificial Gemini.", style='Intense Quote')

    doc.save(output_path)
    print(f"Documento salvo em: {output_path}")

if __name__ == "__main__":
    desktop_path = os.path.join(os.environ['USERPROFILE'], 'Desktop', 'Relatorio_Modulo_Ligacao.docx')
    create_report(desktop_path)
